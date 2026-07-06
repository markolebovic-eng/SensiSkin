from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Base font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

ACCENT = RGBColor(0x8A, 0x4B, 0x5E)  # muted rose, neutral placeholder tone

# Title
title = doc.add_heading('Predlog tema za blog', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
for run in title.runs:
    run.font.color.rgb = ACCENT

sub = doc.add_paragraph('Sensi Skin — jul 2026.')
sub.runs[0].font.size = Pt(13)
sub.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph(
    'U nastavku su predložene teme za naredne blog postove, podeljene u dve grupe. '
    'Za svaku temu je dat kratak opis zašto je zanimljiva i kome je namenjena. '
    'Molimo Vas da izaberete 2–3 teme (iz jedne ili obe grupe) sa kojima želite da krenemo.'
)

# ---------------------------------------------------------------
# GROUP A
# ---------------------------------------------------------------
h1 = doc.add_heading('Grupa 1 — Teme vezane za naše tretmane', level=1)
for run in h1.runs:
    run.font.color.rgb = ACCENT

doc.add_paragraph(
    'Ove teme se naslanjaju na tretmane koje već nudimo (HydraFacial, Mesojet RF, Aura Reality, '
    'lasersku epilaciju) i prirodno vode čitaoca ka zakazivanju.'
)

group_a = [
    ("Nega kože pre i posle tretmana (HydraFacial / Mesojet)",
     "Praktični saveti šta raditi pre i posle tretmana da rezultat traje duže. Prirodno vodi čitaoca ka zakazivanju konsultacije."),
    ("Zašto tretmani koji podstiču kolagen (Mesojet RF) rastu brže od klasičnih zahvata",
     "Objašnjava zašto sve više ljudi bira postepene tretmane poput Mesojet RF umesto naglih zahvata — u skladu sa globalnim trendom."),
    ("HydraFacial kao mesečna rutina, ne jednokratni tretman",
     "Edukuje klijentkinje zašto redovnost (na 6–8 nedelja) daje bolje rezultate od pojedinačnog tretmana — podržava ponovne posete."),
    ("Zašto sve više muškaraca bira lasersku epilaciju",
     "Prati rastući trend muške epilacije i cilja mušku publiku koja tek razmatra ovu uslugu."),
    ("LED terapija kod kuće vs. profesionalni Dermalux tretman",
     "Poredi popularne kućne LED uređaje sa profesionalnim tretmanom u studiju i objašnjava razliku u efikasnosti. "
     "Napomena: postoji već napisan test-draft na sličnu temu (osnovni mehanizam LED terapije) koji čeka odabir — ako izaberete ovu temu, uglovi se mogu spojiti."),
    ("AI aplikacije za analizu kože vs. profesionalna Aura Reality 3D dijagnostika",
     "Poredi popularne AI aplikacije za \"skeniranje\" kože sa stručnom 3D dijagnostikom i ističe prednost stručnog tumačenja. "
     "Napomena: postoji već napisan test-draft na sličnu temu (analiza kože pre tretmana) koji čeka odabir — ako izaberete ovu temu, uglovi se mogu spojiti."),
    ("Retinol i vitamin C 2.0 — zašto isti sastojci sada rade bolje",
     "Objašnjava zašto nove formulacije poznatih sastojaka daju bolje rezultate — prirodno vodi ka preporuci pravih proizvoda iz naše ponude."),
    ("Zaštita kožne barijere — novi fokus u prevenciji perutanja",
     "Nadovezuje se na temu perutanja kože, sa fokusom na prevenciju umesto samo na lečenje simptoma."),
    ("Zašto \"isti tretman za sve\" više ne funkcioniše",
     "Objašnjava zašto je personalizovan pristup bolji od gotovih paketa — podržava naš način rada i pristup klijentkinjama."),
    ("Aura Reality 3D dijagnostika u medijima — zašto se o njoj piše",
     "Iskorišćava to što je Sensi Skin već pominjan u medijima za ovu uslugu i dodatno jača postojeći ugled."),
]

for i, (naslov, opis) in enumerate(group_a, start=1):
    p = doc.add_paragraph()
    run_num = p.add_run(f'{i}. ')
    run_num.bold = True
    run_num.font.color.rgb = ACCENT
    run_title = p.add_run(naslov)
    run_title.bold = True
    p.add_run('\n' + opis).font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(10)

# ---------------------------------------------------------------
# GROUP B
# ---------------------------------------------------------------
h2 = doc.add_heading('Grupa 2 — Opšte teme o nezi kože i lepoti', level=1)
for run in h2.runs:
    run.font.color.rgb = ACCENT

doc.add_paragraph(
    'Ove teme nisu direktno vezane za konkretan tretman koji nudimo, već pokrivaju široko interesovanje '
    'za negu kože i lepotu. Cilj je da privučemo širu publiku i novi saobraćaj na sajt, ne samo ljude koji već traže naše usluge.'
)

group_b = [
    ("Pravilan redosled nanošenja proizvoda za negu kože",
     "Jedna od najtraženijih tema uopšte — svako ko koristi više proizvoda za negu želi da zna kojim redom ih nanosi."),
    ("Skincare mitovi razotkriveni",
     "Popularan i lako deljiv format (mit/istina) koji privlači širu publiku, bez obzira na to da li su ikada bili naši klijenti."),
    ("San, stres i koža",
     "Objašnjava kako svakodnevni stres i loš san utiču na izgled kože — relevantno je svakome, ne samo klijentima."),
    ("Ishrana i koža",
     "Povezuje ishranu sa stanjem kože — evergreen tema koja zanima gotovo svakoga."),
    ("Skin cycling — aktuelan trend nege kože",
     "Aktuelan trend sa društvenih mreža (TikTok/Instagram) — dobra prilika da uhvatimo trenutno interesovanje publike."),
    ("Kako čitati deklaraciju kozmetičkog proizvoda",
     "Praktičan vodič koji pomaže svakom kupcu kozmetike i gradi naš ugled kao stručnog izvora informacija."),
    ("Sezonska promena rutine nege kože (leto/zima)",
     "Klasična, uvek aktuelna tema — svako povremeno menja svoju rutinu u zavisnosti od godišnjeg doba."),
    ("Da li kolagen suplementi zaista deluju na kožu",
     "Aktuelna potrošačka tema o kojoj se dosta priča poslednjih meseci — objektivan, edukativan ugao."),
    ("Prirodna vs. sintetička kozmetika — da li je \"prirodno\" uvek bezbednije",
     "Razbija čest mit da je \"prirodno\" automatski bezbednije ili bolje od sintetičkog — zanimljiv i lako deljiv ugao."),
]

for i, (naslov, opis) in enumerate(group_b, start=1):
    p = doc.add_paragraph()
    run_num = p.add_run(f'{i}. ')
    run_num.bold = True
    run_num.font.color.rgb = ACCENT
    run_title = p.add_run(naslov)
    run_title.bold = True
    p.add_run('\n' + opis).font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(10)

# ---------------------------------------------------------------
# Footer note
# ---------------------------------------------------------------
doc.add_heading('Sledeći korak', level=1)
doc.add_paragraph(
    'Molimo označite ili nam pošaljite nazive 2–3 tema (iz jedne ili obe grupe) koje želite da prve obradimo. '
    'Za izabrane teme pišemo pun tekst za blog, prilagođen našem tonu i stilu.'
)

doc.save(r'C:\Users\Marko\SensiSkin\outputs\sensiskin\blog\predlog-tema-za-blog-2026-07-06.docx')
print('saved')
